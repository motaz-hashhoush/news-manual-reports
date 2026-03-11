import io
import os
import re
from collections import Counter, defaultdict
from datetime import datetime

from docx import Document
from app.timezone import now_palestine
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.config import UPLOAD_DIR

STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
REPORTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "reports")
os.makedirs(REPORTS_DIR, exist_ok=True)

FONT_NAME = "Arial"

ARABIC_MONTHS = {
    1: "يناير", 2: "فبراير", 3: "مارس", 4: "أبريل",
    5: "مايو", 6: "يونيو", 7: "يوليو", 8: "أغسطس",
    9: "سبتمبر", 10: "أكتوبر", 11: "نوفمبر", 12: "ديسمبر",
}

ARABIC_HOURS = {
    1: "الأولى", 2: "الثانية", 3: "الثالثة", 4: "الرابعة",
    5: "الخامسة", 6: "السادسة", 7: "السابعة", 8: "الثامنة",
    9: "التاسعة", 10: "العاشرة", 11: "الحادية عشرة", 12: "الثانية عشرة",
    13: "الثالثة عشرة", 14: "الرابعة عشرة", 15: "الخامسة عشرة",
    16: "السادسة عشرة", 17: "السابعة عشرة", 18: "الثامنة عشرة",
    19: "التاسعة عشرة", 20: "العشرون", 21: "الحادية والعشرون",
    22: "الثانية والعشرون", 23: "الثالثة والعشرون",
    24: "الرابعة والعشرون",
}

DIST_DESCRIPTIONS = {
    "مذيع": "تقديم الأخبار المتجددة والعاجلة والتطورات اللحظية",
    "مراسل": "مداخلات ميدانية حيّة من مواقع مختلفة",
    "ضيف": "لقاءات تحليلية متخصصة",
    "مسؤول": "تصريحات مباشرة لمسؤولين كبار",
    "وول": "مواد تحليلية شاملة ومعمّقة",
    "فيلر": "مواد تفسيرية مصوّرة",
    "تقرير": "تقارير ميدانية مصوّرة",
    "عاجل": "أخبار عاجلة ذات أولوية قصوى",
    "تحليل": "مواد تحليلية",
}


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(
        qn("w:shd"),
        {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex},
    )
    shading.append(shading_elm)


def _style_header_row(row, bg_color="1F4E79"):
    for cell in row.cells:
        _set_cell_shading(cell, bg_color)
        _set_cell_rtl(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = FONT_NAME
                _set_run_rtl_font(run)


def _set_cell_rtl(cell):
    for paragraph in cell.paragraphs:
        pPr = paragraph._element.get_or_add_pPr()
        bidi = pPr.makeelement(qn("w:bidi"), {})
        pPr.append(bidi)


def _set_run_rtl_font(run):
    rPr = run._element.get_or_add_rPr()
    cs = rPr.makeelement(qn("w:rFonts"), {qn("w:cs"): FONT_NAME})
    rPr.append(cs)


def _is_valid_url(text):
    return bool(re.match(r"https?://\S+", text.strip()))


def _append_hyperlink(paragraph, url, display_text):
    """Append a clickable HYPERLINK field to an existing paragraph."""
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f' HYPERLINK "{url}" '
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    run_text = paragraph.add_run(display_text)
    run_text.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    run_text.font.underline = True
    run_text.font.name = FONT_NAME
    _set_run_rtl_font(run_text)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def _add_hyperlink_to_cell(cell, url, display_text):
    """Replace cell content with a clickable hyperlink."""
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    _append_hyperlink(paragraph, url, display_text)


def _set_table_rtl(table):
    """Set table direction to RTL using bidiVisual."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    bidi_visual = OxmlElement("w:bidiVisual")
    tblPr.append(bidi_visual)


def _add_paragraph(doc, text="", bold=False, size=None, alignment=None, justify=False, font=FONT_NAME):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment is not None:
        p.alignment = alignment
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)
    if text:
        run = p.add_run(text)
        run.font.bold = bold
        run.font.name = font
        _set_run_rtl_font(run)
        if size:
            run.font.size = Pt(size)
    return p


def _add_mixed_paragraph(doc, parts, alignment=None, justify=False):
    """Add a paragraph with mixed bold/normal runs.
    parts: list of (text, bold) tuples.
    """
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment is not None:
        p.alignment = alignment
    pPr = p._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn("w:bidi"), {})
    pPr.append(bidi)
    for text, bold in parts:
        run = p.add_run(text)
        run.font.bold = bold
        run.font.name = FONT_NAME
        _set_run_rtl_font(run)
    return p


def _format_arabic_dt(dt):
    hour = dt.hour
    period = "مساءً" if hour >= 12 else "صباحًا"
    display_hour = hour - 12 if hour > 12 else (12 if hour == 0 else hour)
    if hour == 12:
        display_hour = 12
    minute = dt.minute
    month = ARABIC_MONTHS.get(dt.month, str(dt.month))
    time_str = f"{display_hour:02d}:{minute:02d}"
    return f"{time_str} {period}", f"{dt.day} {month} {dt.year}"


def _parse_entry_hour(monitoring_time):
    """Parse monitoring_time like '07:03 PM' into a sortable hour key and display range."""
    if not monitoring_time:
        return None, None
    m = re.match(r"(\d{1,2}):(\d{2})(?::(\d{2}))?\s*(AM|PM|am|pm)?", monitoring_time.strip())
    if not m:
        return None, None
    hour = int(m.group(1))
    ampm = (m.group(4) or "").upper()
    if ampm == "PM" and hour != 12:
        hour += 12
    elif ampm == "AM" and hour == 12:
        hour = 0
    suffix = "AM" if hour < 12 else "PM"
    display_h = hour % 12 or 12
    start = f"{display_h:02d}:00 {suffix}"
    end_h = (hour + 1) % 24
    end_suffix = "AM" if end_h < 12 else "PM"
    end_display = end_h % 12 or 12
    end = f"{end_display:02d}:59 {end_suffix}"
    return hour, f"{start} – {end}"


def _format_entry_timing(entry, include_date=False):
    base_time = getattr(entry, "monitoring_time", None) or "—"
    if not include_date:
        return base_time

    created_at = getattr(entry, "created_at", None)
    if not created_at:
        return base_time

    _, arabic_date = _format_arabic_dt(created_at)
    return f"{base_time} | {arabic_date}"


def _style_table_body(table, font_size=9):
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        for cell in row.cells:
            _set_cell_rtl(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = FONT_NAME
                    _set_run_rtl_font(run)


def generate_docx_report(report_session, entries, breaking_news_count: int = 0) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style_element = style.element
    rPr = style_element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {qn("w:cs"): FONT_NAME})
    rPr.append(rFonts)

    total = len(entries)
    dist_counter = Counter(e.distribution for e in entries if e.distribution)
    type_counter = Counter(e.entry_type for e in entries if e.entry_type)
    program_counter = Counter(e.program for e in entries if e.program)
    published_entries = [e for e in entries if e.publish_link and e.publish_link.strip() and "لم ينشر" not in e.publish_link]
    published = len(published_entries)
    correspondents = [e for e in entries if e.distribution == "مراسل" and e.guest_reporter_name]
    guests = [e for e in entries if e.distribution == "ضيف" and e.guest_reporter_name]
    officials = [e for e in entries if e.distribution == "مسؤول" and e.guest_reporter_name]
    unique_correspondents = Counter(e.guest_reporter_name for e in correspondents)
    unique_guests = Counter(e.guest_reporter_name for e in guests)
    filler_wall_count = dist_counter.get("فيلر", 0) + dist_counter.get("وول", 0)
    report_count = dist_counter.get("تقرير", 0)

    # ── COVER ──
    doc.add_paragraph("")
    _add_paragraph(doc, "تقرير الرصد والتحليل الشامل", bold=True, size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "التغطية الإخبارية المتواصلة — قناة الشرق", bold=True, size=16,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)

    desc = report_session.description or "[يُعبّأ يدويًا: عنوان الحدث / الموضوع الرئيسي]"
    _add_paragraph(doc, desc, bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if report_session.start_at and report_session.deadline_at:
        start_dt = report_session.start_at
        end_dt = report_session.deadline_at
        month = ARABIC_MONTHS.get(start_dt.month, str(start_dt.month))
        if start_dt.day == end_dt.day:
            date_str = f"{start_dt.day} {month} {start_dt.year}"
        else:
            date_str = f"{start_dt.day}- {end_dt.day} {month} {start_dt.year}"
        start_time, _ = _format_arabic_dt(start_dt)
        end_time, _ = _format_arabic_dt(end_dt)
        period_str = f"من الساعة {start_time} وحتى الساعة {end_time}"
        if start_dt.day != end_dt.day:
            period_str += " من اليوم التالي"
        diff_hours = int((end_dt - start_dt).total_seconds() / 3600)
    else:
        date_str = now_palestine().strftime("%Y-%m-%d")
        period_str = "خلال فترة الجلسة"
        diff_hours = getattr(report_session, "duration_hours", 24) or 24

    _add_mixed_paragraph(doc, [("التاريخ", True), (f" {date_str}", False)])
    _add_mixed_paragraph(doc, [("فترة الرصد", True), (": ", True), (f" {period_str}", False)])
    _add_mixed_paragraph(doc, [("إجمالي البث المتواصل", True), (": ", True),
                               (f"نحو {diff_hours} ساعة متواصلة", False)])

    programs_list = " | ".join(program_counter.keys()) if program_counter else "—"
    _add_mixed_paragraph(doc, [("البرامج المشمولة", True), (": ", True), (programs_list, False)])

    doc.add_paragraph("")

    # ── الملخص التنفيذي ──
    _add_paragraph(doc, "الملخص التنفيذي", bold=True, size=10)

    summary = (
        f"يرصد هذا التقرير {total} مادة إخبارية بثّتها قناة الشرق على مدار نحو "
        f"{diff_hours} ساعة متواصلة. "
    )
    if dist_counter:
        parts = []
        for d, c in dist_counter.most_common():
            parts.append(f"{c} مادة {d}")
        summary += "توزعت المواد بين: " + "، ".join(parts) + ". "
    if published:
        summary += f"تم نشر {published} مادة عبر المنصات الرقمية. "
    if breaking_news_count:
        summary += f"رُصد {breaking_news_count} خبر عاجل خلال الفترة."

    _add_paragraph(doc, summary, justify=True)

    _add_paragraph(doc,
        "[يُعبّأ يدويًا: فقرة تحليلية تتناول أبرز محاور التغطية والسياق العام للأحداث]",
        justify=True)

    doc.add_paragraph("")

    # ── المؤشرات الرئيسية ──
    _add_paragraph(doc, "المؤشرات الرئيسية", bold=True, size=10)

    indicators = [
        ("إجمالي المواد المبثوثة", f"{total} مادة"),
        ("إجمالي ساعات البث المتواصل", f"≈ {diff_hours} ساعة"),
        ("عدد مداخلات المراسلين والصحفيين", f"{len(correspondents)} مداخلة"),
        ("عدد المراسلين", f"{len(unique_correspondents)} مراسلاً وصحفياً"),
        ("عدد الضيوف والخبراء", f"{len(unique_guests)}+ خبيراً وضيفاً"),
        ("المواد التحليلية (فيلر + وول)", f"{filler_wall_count} مادة"),
        ("تصريحات كبار المسؤولين", f"{len(officials)} مادة"),
        ("التقارير الميدانية المصوّرة", f"{report_count} تقارير"),
    ]

    t_ind = doc.add_table(rows=1 + len(indicators), cols=2)
    t_ind.style = "Table Grid"
    t_ind.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_rtl(t_ind)
    t_ind.rows[0].cells[0].text = "المؤشر"
    t_ind.rows[0].cells[1].text = "القيمة"
    _style_header_row(t_ind.rows[0])
    for i, (label, value) in enumerate(indicators, 1):
        t_ind.rows[i].cells[0].text = label
        t_ind.rows[i].cells[1].text = value
    _style_table_body(t_ind)

    doc.add_paragraph("")

    # ── توزيع المواد حسب شكل التقديم ──
    _add_paragraph(doc, "توزيع المواد حسب شكل التقديم", bold=True, size=10)

    dist_items = dist_counter.most_common()
    t_dist = doc.add_table(rows=1 + len(dist_items), cols=3)
    t_dist.style = "Table Grid"
    t_dist.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_rtl(t_dist)
    t_dist.rows[0].cells[0].text = "شكل التقديم"
    t_dist.rows[0].cells[1].text = "العدد"
    t_dist.rows[0].cells[2].text = "الوصف الوظيفي"
    _style_header_row(t_dist.rows[0])
    for i, (dist, count) in enumerate(dist_items, 1):
        t_dist.rows[i].cells[0].text = dist
        t_dist.rows[i].cells[1].text = f"{count} مادة"
        t_dist.rows[i].cells[2].text = DIST_DESCRIPTIONS.get(dist, "—")
    _style_table_body(t_dist)

    doc.add_paragraph("")

    # ── توزيع المواد حسب التصنيف الموضوعي ──
    _add_paragraph(doc, "توزيع المواد حسب التصنيف الموضوعي", bold=True, size=10)

    type_items = type_counter.most_common()
    t_type = doc.add_table(rows=1 + len(type_items), cols=3)
    t_type.style = "Table Grid"
    t_type.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_rtl(t_type)
    t_type.rows[0].cells[0].text = "التصنيف"
    t_type.rows[0].cells[1].text = "العدد"
    t_type.rows[0].cells[2].text = "النسبة"
    _style_header_row(t_type.rows[0])
    for i, (etype, count) in enumerate(type_items, 1):
        pct = f"{(count / total * 100):.0f}%" if total else "0%"
        t_type.rows[i].cells[0].text = etype
        t_type.rows[i].cells[1].text = f"≈ {count} مادة"
        t_type.rows[i].cells[2].text = pct
    _style_table_body(t_type)

    doc.add_paragraph("")

    # ── شبكة المراسلين والصحفيين ──
    _add_paragraph(doc, "شبكة المراسلين والصحفيين", bold=True, size=10)

    if unique_correspondents:
        corr_data = []
        for name, count in unique_correspondents.most_common():
            sample = next((e for e in correspondents if e.guest_reporter_name == name), None)
            topic = sample.title if sample and sample.title else "—"            
            corr_data.append((name, count, topic))

        t_corr = doc.add_table(rows=1 + len(corr_data), cols=4)
        t_corr.style = "Table Grid"
        t_corr.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_rtl(t_corr)
        t_corr.rows[0].cells[0].text = "المراسل/الصحفي"
        t_corr.rows[0].cells[1].text = "الموقع"
        t_corr.rows[0].cells[2].text = "عدد المداخلات"
        t_corr.rows[0].cells[3].text = "الملف الرئيسي"
        _style_header_row(t_corr.rows[0])
        for i, (name, count, topic) in enumerate(corr_data, 1):
            t_corr.rows[i].cells[0].text = name
            t_corr.rows[i].cells[1].text = "—"
            t_corr.rows[i].cells[2].text = str(count)
            t_corr.rows[i].cells[3].text = topic
        _style_table_body(t_corr)
    else:
        _add_paragraph(doc, "لا يوجد مراسلون في هذه الفترة", justify=True)

    doc.add_paragraph("")

    # ── أبرز الضيوف والخبراء ──
    _add_paragraph(doc, "أبرز الضيوف والخبراء", bold=True, size=10)

    if guests:
        seen_guests = {}
        for e in guests:
            if e.guest_reporter_name not in seen_guests:
                seen_guests[e.guest_reporter_name] = e.title if e.title else "—"
        t_guest = doc.add_table(rows=1 + len(seen_guests), cols=3)
        t_guest.style = "Table Grid"
        t_guest.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_rtl(t_guest)
        t_guest.rows[0].cells[0].text = "الضيف"
        t_guest.rows[0].cells[1].text = "الصفة"
        t_guest.rows[0].cells[2].text = "الموضوع الرئيسي"
        _style_header_row(t_guest.rows[0])
        for i, (name, topic) in enumerate(seen_guests.items(), 1):
            t_guest.rows[i].cells[0].text = name
            t_guest.rows[i].cells[1].text = "—"
            t_guest.rows[i].cells[2].text = topic
        _style_table_body(t_guest)
    else:
        _add_paragraph(doc, "لا يوجد ضيوف في هذه الفترة", justify=True)

    doc.add_paragraph("")

    # ── أبرز تصريحات المسؤولين المبثوثة ──
    _add_paragraph(doc, "أبرز تصريحات المسؤولين المبثوثة", bold=True, size=10)

    if officials:
        t_off = doc.add_table(rows=1 + len(officials), cols=3)
        t_off.style = "Table Grid"
        t_off.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_rtl(t_off)
        t_off.rows[0].cells[0].text = "المسؤول"
        t_off.rows[0].cells[1].text = "الصفة"
        t_off.rows[0].cells[2].text = "أبرز ما جاء في التصريح"
        _style_header_row(t_off.rows[0])
        for i, e in enumerate(officials, 1):
            t_off.rows[i].cells[0].text = e.guest_reporter_name or "—"
            t_off.rows[i].cells[1].text = "—"
            t_off.rows[i].cells[2].text = e.title or "—"
        _style_table_body(t_off)
    else:
        _add_paragraph(doc, "لا يوجد تصريحات مسؤولين في هذه الفترة", justify=True)

    doc.add_paragraph("")

    is_custom = bool(getattr(report_session, "is_custom_report", False))

    # ── العرض التفصيلي ساعة بساعة ──
    _add_paragraph(doc, "العرض التفصيلي ساعة بساعة", bold=True, size=10)
    doc.add_paragraph("")

    if is_custom:
        # Group by (date, hour) so multi-day reports get separate day sections
        date_hourly = defaultdict(lambda: defaultdict(list))
        for e in entries:
            h, _ = _parse_entry_hour(e.monitoring_time)
            entry_date = getattr(e, "created_at", None)
            date_key = entry_date.date() if entry_date else None
            date_hourly[date_key][h if h is not None else -1].append(e)

        global_idx = 1
        for date_key in sorted(date_hourly.keys(), key=lambda d: d or datetime.min.date()):
            hours_map = date_hourly[date_key]
            if date_key:
                month = ARABIC_MONTHS.get(date_key.month, str(date_key.month))
                date_label = f"{date_key.day} {month} {date_key.year}"
            else:
                date_label = "تاريخ غير محدد"
            _add_paragraph(doc, f"— {date_label} —", bold=True, size=11)
            doc.add_paragraph("")

            hour_num = 0
            for h_key in sorted(hours_map.keys()):
                hour_entries = hours_map[h_key]
                hour_num += 1

                if h_key >= 0:
                    _, time_range = _parse_entry_hour(hour_entries[0].monitoring_time)
                else:
                    time_range = "—"

                hour_label = ARABIC_HOURS.get(hour_num, f"الساعة {hour_num}")
                _add_paragraph(doc, f"الساعة {hour_label} | {time_range}", bold=True, size=10)

                _add_mixed_paragraph(doc, [
                    ("تحليل الساعة", True), (": ", True),
                    (f"{len(hour_entries)} ", False),
                    ("مادة ", False),
                    ("— [يُعبّأ يدويًا: تحليل موجز لأبرز أحداث هذه الساعة]", False),
                ])
                doc.add_paragraph("")

                t_hour = doc.add_table(rows=1 + len(hour_entries), cols=6)
                t_hour.style = "Table Grid"
                t_hour.alignment = WD_TABLE_ALIGNMENT.CENTER
                _set_table_rtl(t_hour)
                h_row = t_hour.rows[0]
                for j, h_text in enumerate(["#", "التوقيت", "المادة", "النوع", "التوزيع", "الضيف/المراسل"]):
                    h_row.cells[j].text = h_text
                _style_header_row(h_row)

                for i, entry in enumerate(hour_entries):
                    row = t_hour.rows[i + 1]
                    link = entry.publish_link or ""
                    has_link = link.strip() and "لم ينشر" not in link
                    guest_cell = entry.guest_reporter_name or "—"
                    if has_link and _is_valid_url(link):
                        p = row.cells[5].paragraphs[0]
                        p.clear()
                        if guest_cell != "—":
                            p.add_run(guest_cell + " | ")
                        _append_hyperlink(p, link.strip(), "رابط")
                    elif has_link:
                        row.cells[5].text = (guest_cell + " | " + link.strip()) if guest_cell != "—" else link.strip()
                    else:
                        row.cells[5].text = guest_cell
                    row.cells[4].text = entry.distribution or "—"
                    row.cells[3].text = entry.entry_type or "—"
                    row.cells[2].text = entry.title or "—"      
                    row.cells[1].text = _format_entry_timing(entry, include_date=True)
                    row.cells[0].text = str(global_idx)
                    global_idx += 1

                _style_table_body(t_hour, font_size=8)
                doc.add_paragraph("")
    else:
        hourly = defaultdict(list)
        for e in entries:
            h, _ = _parse_entry_hour(e.monitoring_time)
            hourly[h if h is not None else -1].append(e)

        global_idx = 1
        hour_num = 0
        for h_key in sorted(hourly.keys()):
            hour_entries = hourly[h_key]
            hour_num += 1

            if h_key >= 0:
                _, time_range = _parse_entry_hour(hour_entries[0].monitoring_time)
            else:
                time_range = "—"

            hour_label = ARABIC_HOURS.get(hour_num, f"الساعة {hour_num}")
            _add_paragraph(doc, f"الساعة {hour_label} | {time_range}", bold=True, size=10)

            _add_mixed_paragraph(doc, [
                ("تحليل الساعة", True), (": ", True),
                (f"{len(hour_entries)} ", False),
                ("مادة ", False),
                ("— [يُعبّأ يدويًا: تحليل موجز لأبرز أحداث هذه الساعة]", False),
            ])

            doc.add_paragraph("")

            t_hour = doc.add_table(rows=1 + len(hour_entries), cols=6)
            t_hour.style = "Table Grid"
            t_hour.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_rtl(t_hour)

            h_row = t_hour.rows[0]
            for j, h_text in enumerate(["#", "التوقيت", "المادة", "النوع", "التوزيع", "الضيف/المراسل"]):
                h_row.cells[j].text = h_text
            _style_header_row(h_row)

            for i, entry in enumerate(hour_entries):
                row = t_hour.rows[i + 1]
                link = entry.publish_link or ""
                has_link = link.strip() and "لم ينشر" not in link
                guest_cell = entry.guest_reporter_name or "—"
                if has_link and _is_valid_url(link):
                    p = row.cells[5].paragraphs[0]
                    p.clear()
                    if guest_cell != "—":
                        p.add_run(guest_cell + " | ")
                    _append_hyperlink(p, link.strip(), "رابط")
                elif has_link:
                    row.cells[5].text = (guest_cell + " | " + link.strip()) if guest_cell != "—" else link.strip()
                else:
                    row.cells[5].text = guest_cell
                row.cells[4].text = entry.distribution or "—"
                row.cells[3].text = entry.entry_type or "—"
                row.cells[2].text = entry.title or "—"                
                row.cells[1].text = entry.monitoring_time or "—"
                row.cells[0].text = str(global_idx)
                global_idx += 1

            _style_table_body(t_hour, font_size=8)
            doc.add_paragraph("")

    # ── نماذج من المواد المنشورة رقمياً ──
    _add_paragraph(doc, "نماذج من المواد المنشورة رقمياً", bold=True, size=10)

    if published_entries:
        t_pub = doc.add_table(rows=1 + len(published_entries), cols=3)
        t_pub.style = "Table Grid"
        t_pub.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_rtl(t_pub)
        t_pub.rows[0].cells[0].text = "التوقيت"
        t_pub.rows[0].cells[1].text = "المادة"
        t_pub.rows[0].cells[2].text = "الرابط"
        _style_header_row(t_pub.rows[0])
        for i, e in enumerate(published_entries, 1):
            link = e.publish_link or ""
            if _is_valid_url(link):
                if "facebook" in link.lower() or "fb.com" in link.lower():
                    display_text = "Facebook"
                elif "twitter" in link.lower() or "x.com" in link.lower():
                    display_text = "X/Twitter"
                else:
                    display_text = "رابط"
                _add_hyperlink_to_cell(t_pub.rows[i].cells[2], link.strip(), display_text)
            else:
                t_pub.rows[i].cells[2].text = link.strip() if link.strip() else "—"
            t_pub.rows[i].cells[1].text = (e.title or "—")
            t_pub.rows[i].cells[0].text = _format_entry_timing(e, include_date=is_custom)
        _style_table_body(t_pub)
    else:
        _add_paragraph(doc, "لا توجد مواد منشورة رقمياً في هذه الفترة", justify=True)
    doc.add_paragraph("")

    # ── Screenshots ──
    screenshot_entries = [e for e in entries if e.screenshot_path or e.screenshot_data]
    if screenshot_entries:
        _add_paragraph(doc, "نماذج (سكرين شوت)", bold=True, size=10)
        for sc_entry in screenshot_entries:
            added = False
            if sc_entry.screenshot_path:
                filename = os.path.basename(sc_entry.screenshot_path)
                abs_path = os.path.join(UPLOAD_DIR, filename)
                if not os.path.exists(abs_path):
                    abs_path = os.path.join(STATIC_DIR, sc_entry.screenshot_path.lstrip("/").removeprefix("static/"))
                if os.path.exists(abs_path):
                    try:
                        doc.add_picture(abs_path, width=Inches(5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        added = True
                    except Exception:
                        pass
            if not added and sc_entry.screenshot_data:
                try:
                    image_stream = io.BytesIO(sc_entry.screenshot_data)
                    doc.add_picture(image_stream, width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    added = True
                except Exception:
                    pass
            if not added:
                _add_paragraph(doc, f"[صورة غير متوفرة: {sc_entry.title}]")
        doc.add_paragraph("")

    # ── الخلاصات والملاحظات التحريرية ──
    _add_paragraph(doc, "الخلاصات والملاحظات التحريرية", bold=True, size=10)

    _add_mixed_paragraph(doc, [
        ("أولاً — الإيقاع والكثافة", True), (": ", True),
        ("[يُعبّأ يدويًا: ملاحظات حول إيقاع التغطية ومعدل المواد في الساعة]", False),
    ], justify=True)

    _add_mixed_paragraph(doc, [
        ("ثانياً — شبكة المراسلين", True), (": ", True),
        (f"{len(unique_correspondents)} مراسلاً قدّموا {len(correspondents)} مداخلة. ", False),
        ("[يُعبّأ يدويًا: ملاحظات حول التوزيع الجغرافي وأداء المراسلين]", False),
    ], justify=True)

    _add_mixed_paragraph(doc, [
        ("ثالثاً — التنوّع التحليلي", True), (": ", True),
        (f"أكثر من {len(unique_guests)} ضيفاً. ", False),
        ("[يُعبّأ يدويًا: ملاحظات حول تنوّع الضيوف والتخصصات]", False),
    ], justify=True)

    _add_mixed_paragraph(doc, [
        ("رابعاً — التكامل بين الأخبار والاقتصاد", True), (": ", True),
        ("[يُعبّأ يدويًا: ملاحظات حول التكامل بين الأخبار والبرامج الاقتصادية]", False),
    ], justify=True)

    # ── Save ──
    timestamp = now_palestine().strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(c if c.isalnum() or c in (" ", "-", "_") else "_" for c in report_session.name).strip()
    filename = f"{safe_name}_{timestamp}.docx"
    file_path = os.path.join(REPORTS_DIR, filename)
    doc.save(file_path)

    return file_path


class _CustomReportSession:
    def __init__(self, start_at, end_at, custom_name=None):
        self.name = custom_name or "تقرير مخصص حسب فترة زمنية"
        self.description = "التطورات الإخبارية"
        self.start_at = start_at
        self.deadline_at = end_at
        self.duration_hours = max(1, int((end_at - start_at).total_seconds() / 3600))
        self.is_custom_report = True


def generate_custom_docx_report(start_dt, end_dt, entries, breaking_news_count: int = 0, report_name: str = None) -> str:
    fake_session = _CustomReportSession(start_dt, end_dt, report_name)
    return generate_docx_report(
        report_session=fake_session,
        entries=entries,
        breaking_news_count=breaking_news_count,
    )
